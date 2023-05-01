                       Работа с документами PDF и Word

================ PDF документы =======================================

pip install --user PyPDF2
 
>>> import PyPDF2
>>> pdfFileObj = open('meetingminutes.pdf', 'rb')
>>> pdfReader = PyPDF2.PdfReader(pdfFileObj)

>>> len(pdfReader.pages)
19
>>> pageObj = pdfReader.pages[0]

>>> pageObj.extract_text()
'OOFFFFIICCIIAALL  BBOOAARRDD  MMIINNUUTTEESS  \n \nMeeting of March 7 ,
2014  \n \n \n \n  \n \n  \n \n   \nThe Board of Elementary and Secondary Education
shall provide leadership and \ncreate policies for education that
expand opportunities for children, empower \nfamilies and
communities, and advance Louisiana in an increasingly
\ncompetitive glob al market.  BOARD  \nof \nELEMENTARY  \nand  \nSECONDARY
\nEDUCATION  \n '

>>> pdfFileObj.close()



Дешифровка PDF-документов:
--------------------------
>>> import PyPDF2
>>> pdfReader = PyPDF2.PdfReader(open('encrypted.pdf', 'rb'))
>>> pdfReader.is_encrypted
True
>>> pdfReader.decrypt('rosebud')
<PasswordType.OWNER_PASSWORD: 2>
>>> pageObj = pdfReader.pages[0]
>>> pageObj.extract_text()



Создание PDF-документов:
------------------------
Копирование страниц

import PyPDF2
pdf1File = open('meetingminutes.pdf', 'rb')
pdf2File = open('meetingminutes2.pdf', 'rb')

pdf1Reader = PyPDF2.PdfReader(pdf1File)
pdf2Reader = PyPDF2.PdfReader(pdf2File)

pdfWriter = PyPDF2.PdfWriter()

for pageNum in range(len(pdf1Reader.pages)):
    pageObj = pdf1Reader.pages[pageNum]
    pdfWriter.add_page(pageObj)

for pageNum in range(len(pdf2Reader.pages)):
    pageObj = pdf2Reader.pages[pageNum]
    pdfWriter.add_page(pageObj)

pdfOutputFile = open('combinedminutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()    



Поворот страниц:
----------------
import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfReader(minutesFile)
page = pdfReader.pages[0]
page.rotate(90)
--Опущено--

pdfWriter = PyPDF2.PdfWriter()
pdfWriter.add_page(page)
resultPdfFile = open('rotatedPage.pdf', 'wb')
pdfWriter.write(resultPdfFile)

resultPdfFile.close()
minutesFile.close()



Наложение страниц:
------------------
import PyPDF2

minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfReader(minutesFile)
minutesFirstPage = pdfReader.pages[0]

pdfWatermarkReader = PyPDF2.PdfReader(open('watermark.pdf', 'rb'))

minutesFirstPage.merge_page(pdfWatermarkReader.pages[0])

pdfWriter = PyPDF2.PdfWriter()
pdfWriter.add_page(minutesFirstPage)

for pageNum in range(1, len(pdfReader.pages)):
    pageObj = pdfReader.pages[pageNum]
    pdfWriter.add_page(pageObj)

resultPdfFile = open('watermarkedCover.pdf', 'wb')
pdfWriter.write(resultPdfFile)

minutesFile.close()
resultPdfFile.close()



Шифрование PDF документов:
--------------------------

import PyPDF2
pdfFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfReader(pdfFile)
pdfWriter = PyPDF2.PdfWriter()
for pageNum in range(len(pdfReader.pages)):
    pdfWriter.add_page(pdfReader.pages[pageNum])

pdfWriter.encrypt('swordfish')
resultPdf = open('encryptedminutes.pdf', 'wb')
pdfWriter.write(resultPdf)

resultPdf.close()




====== Проект: обьединение выбранных страниц из многих PDF документов =======================================

#! python3
# combinePdfs.py - обьединяет все PDF-файлы, находящиеся в текущем каталоге в единый PDF-документ

import PyPDF2, os

# Получение списка всех PDF-файлов и сортировка
pdfFiles = []
for filename in os.listdir('.'):
    if filename.endswith('.pdf'):
        pdfFiles.append(filename)
pdfFiles.sort(key=str.lower)      # Сортировка в алфавитном порядке

pdfWriter = PyPDF2.PdfWriter()    # Создаём объект для хранения объединённых PDF-файлов

# Организация цикла по всем PDF-файлам
for filename in pdfFiles:
    pdfFileObj = open(filename, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFileObj)
    
    # Организация цикла по всем страницам за исключением первой, которые добавляются в результ. документ
    for pageNum in range(1, len(pdfReader.pages)):
        pageObj = pdfReader.pages[pageNum]
        pdfWriter.add_page(pageObj)

# Сохранение результирующего PDF-документа в файле
pdfOutput = open('allminutes.pdf', 'wb')
pdfWriter.write(pdfOutput)
print('Done')
pdfOutput.close()




========================= Документы Word ===========================================

pip install python-docx


Чтение документов:
------------------
import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)
7
doc.paragraphs[0].text
'Document Title'
doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)
5
doc.paragraphs[1].runs[0].text
'A plain paragraph with'
doc.paragraphs[1].runs[1].text
' some '
doc.paragraphs[1].runs[2].text
'bold'
doc.paragraphs[1].runs[3].text
' and some '
doc.paragraphs[1].runs[4].text
'italic'



Получение всего текста из файла .docx 
-------------------------------------
Пишем универсальную функцию которую потом можно будет импортировать в программу.
#! python3
# readDocx.py - Получение всего текста из файла

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for i in doc.paragraphs:
        fullText.append(i.text)
    return '\n'.join(fullText)  # Все строки, содержащиеся в списке fullText
                                # обьединяются с использованием новой строки \n

***

>>> import readDocx
>>> print(readDocx.getText('demo.docx'))
Document Title
A plain paragraph with some bold and some italic
Heading, level 1
Intense quote
first item in unordered list
first item in ordered list



Атрибуты обьекта Run:
---------------------

>>> import docx
>>> doc = docx.Document('demo.docx')
>>> doc.paragraphs[0].text
'Document Title'

>>> doc.paragraphs[0].style
_ParagraphStyle('Title') id: 2166650200144
>>> doc.paragraphs[0].style = 'Normal'

>>> doc.paragraphs[1].text
'A plain paragraph with some bold and some italic'
>>> (doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
('A plain paragraph with', ' some ', 'bold', ' and some ')

>>> doc.paragraphs[1].runs[0].style = 'QuoteChar'
>>> doc.paragraphs[1].runs[1].underline = True
>>> doc.paragraphs[1].runs[3].underline = True
>>> doc.save('restyled.docx')


Запись документов Word:
----------------------

import docx
doc = docx.Document()
doc.add_paragraph('Здравствуй, мир!')
<docx.text.paragraph.Paragraph object at 0x0000021D36140E90>
doc.save('helloworld.docx')


import docx
doc = docx.Document()
doc.add_paragraph('Hello, world!') # Для оформления doc.add_paragraph('Hello world!', 'Title')  
<docx.text.paragraph.Paragraph object at 0x0000021D35E38A90>

paraObj = doc.add_paragraph('Second paragraph.')
paraObj2 =  doc.add_paragraph('Thirty paragraph')
paraObj.add_run(' This text added into Second paragraph.')
<docx.text.run.Run object at 0x0000021D36140E90>
doc.save('multipleParagraphs.docx')

***Результат добавления:   
Hello, world!
Second paragraph. This text added into Second paragraph.
Thirty paragraph



Добавление заголовков(0-4):
--------------------------
Чем больше число, тем меньше заголовок

doc = docx.Document()
doc.add_heading('Header 0', 0)
<docx.text.paragraph.Paragraph object at 0x0000021D361B7710>
doc.add_heading('Header 1', 1)
<docx.text.paragraph.Paragraph object at 0x0000021D357E7FD0>
doc.add_heading('Header 2', 2)
<docx.text.paragraph.Paragraph object at 0x0000021D358AED10>
doc.add_heading('Header 3', 3)
<docx.text.paragraph.Paragraph object at 0x0000021D35F351D0>
doc.add_heading('Header 4', 4)
<docx.text.paragraph.Paragraph object at 0x0000021D3337D450>
doc.save('headings.docx')



Разрыв строки:
--------------
doc = docx.Document()

doc.add_paragraph(' Text on first page')
<docx.text.paragraph.Paragraph object at 0x0000021D35F17D90>

doc.paragraphs[0].runs[0].add_break()

doc.add_paragraph(' Continue')
<docx.text.paragraph.Paragraph object at 0x0000021D35EF2110>

doc.save('Continue.docx')



Добавление страниц:
-------------------
doc = docx.Document()
doc.add_paragraph(' Text on first page')
<docx.text.paragraph.Paragraph object at 0x0000021D3618F450>

doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.add_paragraph(' Text on second page')
<docx.text.paragraph.Paragraph object at 0x0000021D33BBE690>

doc.save('twoPage.docx')



Добавление изображений:
-----------------------
Добавляем в конец уже созданного файла twoPage изображение размерами 1дюйм 4см

doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
<docx.shape.InlineShape object at 0x0000021D33582050>
doc.save('twoPage.docx')

# Необязательные аргументы
width=docx.shared.Inches(1)
height=docx.shared.Cm(4)


===== Создание документов PDF на основе документов Word ==========================

# Этот сценарий выполняется только в Windows, при условии что установлен Microsoft Word
# pip install pywin32

import win32com.client
import docx
wordFilename = 'Some_document.docx'
pdfFilename = 'Some_document.pdf'

doc = docx.Document()
# Здесь должен быть код для создания документа Word
doc.save(wordFilename)

wbFormatPDF = 17 # Числовой код Word для документов PDF
wordObj = win32com.client.Dispatch('Word.Application')

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wbFormatPDF)
docObj.Close()
wordObj.Quit()










